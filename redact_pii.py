#!/usr/bin/env python3
"""PII redaction tool for the supplied Red Herring Prospectus.

Usage:
    python redact_pii.py input.pdf output.docx

The detector is deliberately dependency-light: regexes handle structured PII,
while document-specific gazetteers handle person/company names that are hard
to infer reliably from a legal prospectus without an NER model.
"""

from __future__ import annotations

import argparse
import hashlib
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Callable, Iterable

from docx import Document
from docx.enum.text import WD_BREAK
from pypdf import PdfReader


# ---------------------------------------------------------------------------
# Document-specific gazetteers. These are seed entities found in the supplied
# prospectus. Adding another PII type only requires a new detector + fake maker.
# ---------------------------------------------------------------------------
PERSON_NAMES = [
    "Kushal Subbayya Hegde", "Pushpa Kushal Hegde", "Rajesh Kushal Hegde",
    "Rohit Kushal Hegde", "Rakhi Girija Shetty", "Kushal Hegde",
    "Pushpa Hegde", "Rajesh Hegde", "Rohit Hegde", "Sarthak Malvadkar",
    "Dinesh Hirachand Munot", "Ajay Shriram Patil", "Ram Kumar Tiwari",
    "Indu Jacob", "Pratik Bunglow", "Maithili Rajesh Hegde",
    "Katyayani Balasubramanian", "Rupal K. Sancheti", "Salil Ajay Bhargava",
    "Jabeen Ajay Menon", "Ajay Menon", "Sunil Nagayya Shetty",
    "Lalit Muljibhai Sarvaiya", "Lokesh Shah", "Soumavo Sarkar",
    "Kishan Rastogi", "Abhijit Diwan", "Prakash Boricha",
    "Shanti Gopalkrishnan", "Parag Pansare", "Eric Bacha", "Sachin Gawade",
    "Pravin Teli", "Siddharth Jadhav", "Tushar Gavankar", "Tushar Wakhele",
    "Cherag Gyara", "Manisha Shukla", "Ashish Mathew Pulloor", "Anand Soni",
    "Hitesh Ramani", "Chitra Raste", "Sharmila Joshi", "Sandesh Bhagwat",
    "Amod Joshi", "Ganesh Prasad", "Karunakar Hegde", "Karunakar Bhandary",
    "Karunakar N. Bhandary",
]

# Only legal-entity/company-like names are targeted here; regulators and
# generic phrases such as "Stock Exchanges" are intentionally excluded.
COMPANY_SUFFIX_RE = re.compile(
    r"\b(?:[A-Z][A-Za-z0-9&.'()-]*\s+){0,8}"
    r"(?:Limited|Private Limited|LLP|Ltd\.?|Inc\.?|Corporation)\b"
)

# Additional entity names whose suffix can be lost by PDF text extraction.
COMPANY_NAMES = [
    "KSH International Limited", "Waterloo Industrial Park VI Private Limited",
    "Kirtane & Pandit LLP", "Nuvama Wealth Management Limited",
    "ICICI Securities Limited", "MUFG Intime India Private Limited",
    "HDFC Bank Limited", "ICICI Bank Limited", "Citibank N.A.",
    "Export-Import Bank of India", "IndusInd Bank Limited", "The Federal Bank Limited",
    "Bajaj Finance Limited", "CARE Analytics and Advisory Private Limited",
    "KSH Distriparks Private Limited", "KSH Project Management Services Private Limited",
    "KSH Infra Park 5 Private Limited", "KSH Infra Park VI Private Limited",
    "KSH Integrated Logistics Private Limited", "Kushal Motors and Electricals Private Limited",
    "Waterloo Motors Private Limited", "Waterloo Industrial Park I Private Limited",
    "Waterloo Industrial Park II Private Limited", "Waterloo Industrial Park III Private Limited",
    "Waterloo Industrial Park IV Private Limited", "Waterloo Industrial Park V Private Limited",
    "Waterloo Industrial Park VIII Private Limited", "Waterloo Industrial Park IX Private Limited",
    "Waterloo Industrial Park IX B Private Limited", "Parijat Foundation",
    "KSH Infra Park 4 Private Limited", "Bhandary Metal Extrusion Private Limited",
    "Polycom Associates", "Savli Copper Products Private Limited", "Union Copper Rod LLC",
    "Vedanta Limited Sterlite Copper", "Georgia Transformer Corporation",
    "Virginia Transformer Corporation", "Nidec Industrial Automation India Private Limited",
    "Transformers & Rectifiers (India) Limited", "Bharat Bijlee Limited",
    "CG Power and Industrial Solutions Limited", "Al-Ahleia Switchgear Co.",
    "Emirates Transformer & Switchgear Limited", "Malabar India Fund Limited",
    # Promoter Family Trusts
    "Dhaulagiri Family Trust", "Everest Family Trust", "Makalu Family Trust",
    "Broad Family Trust", "Annapurna Family Trust", "Kanchenjunga Family Trust",
]

EMAIL_RE = re.compile(r"\b[A-Z0-9._%+\-]+@[A-Z0-9.\-]+\.[A-Z]{2,}\b", re.I)
SSN_RE = re.compile(r"\b\d{3}-\d{2}-\d{4}\b")
IP_RE = re.compile(r"\b(?:\d{1,3}\.){3}\d{1,3}\b")
CC_RE = re.compile(r"\b(?:\d[ -]*?){13,19}\b")
DOB_RE = re.compile(
    r"(?i)\b(?:date\s+of\s+birth|dob|born\s+on)\s*[:\-]?\s*"
    r"(?:\d{1,2}[/-]\d{1,2}[/-]\d{2,4}|\d{1,2}\s+[A-Za-z]+\s+\d{4}|[A-Za-z]+\s+\d{1,2},\s*\d{4})"
)
PHONE_RE = re.compile(
    r"(?<!\d)(?:\+\s*)?\d{1,3}[\s().-]*(?:\d[\s().-]*){9,12}\d(?!\d)"
)
PIN_RE = re.compile(r"(?<!\d)\d{3}\s?\d{3}(?!\d)")

ADDRESS_KEYWORDS = re.compile(
    r"(?i)\b(?:road|street|marg|lane|village|taluka|district|plot|tower|floor|block|"
    r"building|campus|park|complex|house|apartment|society|nagar|chambers|bunglow|"
    r"bungalow|industrial area|phase|sector|farm|facility|colony)\b"
)
ADDRESS_LABEL_RE = re.compile(
    r"(?i)\b(?:registered office|corporate office|address(?: of the roC)?|"
    r"registered office of our company)\b[^.]{10,260}?\b\d{6}\b"
)

@dataclass(frozen=True)
class Span:
    start: int
    end: int
    pii_type: str


def normalize(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def luhn_ok(value: str) -> bool:
    digits = [int(c) for c in value if c.isdigit()]
    if not 13 <= len(digits) <= 19:
        return False
    total = 0
    parity = len(digits) % 2
    for i, d in enumerate(digits):
        if i % 2 == parity:
            d *= 2
            if d > 9:
                d -= 9
        total += d
    return total % 10 == 0


def add_regex_spans(text: str, regex: re.Pattern, pii_type: str,
                    validator: Callable[[str], bool] | None = None) -> list[Span]:
    out = []
    for m in regex.finditer(text):
        value = m.group(0)
        if validator is None or validator(value):
            out.append(Span(m.start(), m.end(), pii_type))
    return out


def detect_people(text: str) -> list[Span]:
    spans = []
    # Longest-first avoids matching "Rajesh Hegde" inside the full name.
    for name in sorted(PERSON_NAMES, key=len, reverse=True):
        pattern = re.compile(r"(?<![A-Za-z])" + re.escape(name) + r"(?![A-Za-z])", re.I)
        spans.extend(Span(m.start(), m.end(), "PERSON") for m in pattern.finditer(text))
    return spans


def detect_companies(text: str) -> list[Span]:
    spans = []
    names = sorted(set(COMPANY_NAMES), key=len, reverse=True)
    for name in names:
        pattern = re.compile(r"(?<![A-Za-z])" + re.escape(name) + r"(?![A-Za-z])", re.I)
        spans.extend(Span(m.start(), m.end(), "COMPANY") for m in pattern.finditer(text))
    for m in COMPANY_SUFFIX_RE.finditer(text):
        candidate = m.group(0).strip()
        if len(candidate.split()) >= 2:
            spans.append(Span(m.start(), m.end(), "COMPANY"))
    return spans


def detect_addresses(text: str) -> list[Span]:
    spans = []
    for m in ADDRESS_LABEL_RE.finditer(text):
        spans.append(Span(m.start(), m.end(), "ADDRESS"))

    # Catch addresses ending in an Indian PIN. We anchor on a street/building
    # number or an address keyword, then include the run up to the PIN.
    previous_pin_end = 0
    for pin in PIN_RE.finditer(text):
        start = max(previous_pin_end, pin.start() - 220)
        window = text[start:pin.end()]
        candidates = []
        for m in re.finditer(r"(?i)\b(?:\d{1,4}(?:/\d{1,4})+|\d{1,4},|PCNTDA|Gat No\.?|Plot No\.?|S\.?\s?no\.?|CTS No\.?|Flat No\.?|\d{3,4}-\d{3,4})\b", window):
            candidates.append(m.start())
        for m in ADDRESS_KEYWORDS.finditer(window):
            candidates.append(m.start())
        if candidates:
            s = min(candidates)
            tail = window[s:]
            if ADDRESS_KEYWORDS.search(tail) or re.search(r"\d{1,4}(?:/\d{1,4})+", tail):
                spans.append(Span(start + s, pin.end(), "ADDRESS"))
        previous_pin_end = pin.end()
    return spans


def detect_pii(text: str) -> list[Span]:
    spans = []
    spans += detect_people(text)
    spans += detect_companies(text)
    spans += add_regex_spans(text, EMAIL_RE, "EMAIL")
    spans += add_regex_spans(text, SSN_RE, "SSN")
    spans += add_regex_spans(text, IP_RE, "IP_ADDRESS",
                             lambda v: all(int(x) <= 255 for x in v.split('.')))
    spans += add_regex_spans(text, DOB_RE, "DOB")
    spans += add_regex_spans(text, CC_RE, "CREDIT_CARD", luhn_ok)

    # Phone numbers: prioritize explicit phone labels and +country-code forms.
    for m in PHONE_RE.finditer(text):
        raw = m.group(0)
        digits = re.sub(r"\D", "", raw)
        if not 10 <= len(digits) <= 13:
            continue
        before = text[max(0, m.start()-35):m.start()].lower()
        if raw.strip().startswith('+') or any(k in before for k in ('telephone', 'phone', 'mobile', 'tel')):
            spans.append(Span(m.start(), m.end(), "PHONE"))

    spans += detect_addresses(text)

    return merge_spans(spans)


def merge_spans(spans: Iterable[Span]) -> list[Span]:
    # Prioritize specific types (PERSON, COMPANY, EMAIL, PHONE, SSN, CC, DOB, IP)
    # over generic ADDRESS spans so addresses don't swallow specific entity names.
    type_priority = {
        "EMAIL": 1, "PHONE": 1, "SSN": 1, "CREDIT_CARD": 1, "IP_ADDRESS": 1, "DOB": 1,
        "PERSON": 2, "COMPANY": 2, "ADDRESS": 3
    }
    
    ordered = sorted(spans, key=lambda s: (s.start, type_priority.get(s.pii_type, 4), -(s.end-s.start)))
    kept: list[Span] = []
    for s in ordered:
        if not kept or s.start >= kept[-1].end:
            kept.append(s)
        elif type_priority.get(s.pii_type, 4) < type_priority.get(kept[-1].pii_type, 4):
            # Higher priority type (e.g. PERSON vs ADDRESS overlap): keep higher priority
            if s.end > kept[-1].end:
                # Truncate kept address to start after s.end if applicable
                kept[-1] = s
            else:
                kept[-1] = s
        elif s.end > kept[-1].end and type_priority.get(s.pii_type, 4) == type_priority.get(kept[-1].pii_type, 4):
            kept[-1] = s
    return sorted(kept, key=lambda s: s.start)



# Realistic fake-data pools, matching the assignment's spec format.
_FAKE_PERSON = [
    "John Doe", "Jane Smith", "Peter Parker", "Alex Morgan", "Taylor Brown",
    "Jordan Lee", "Sam Wilson", "Chris Blake", "Morgan Reed", "Casey Quinn",
    "Riley Adams", "Drew Campbell", "Avery Hart", "Jamie Stone", "Skyler West",
    "Logan Grey", "Dakota Mills", "Reese Palmer", "Hayden Brooks", "Finley Shaw",
    "Rowan Clarke", "Ellis Grant", "Kendall Price", "Emery Fox", "Harper Kim",
    "Blake Turner", "Quinn Nash", "Parker Lane", "Sage Porter", "Emerson Cole",
    "Phoenix Day", "Remy Cross", "Devon Ray", "Cameron Bell", "Oakley Wood",
    "Kai Rivera", "Jesse Hale", "Shawn Marsh", "Max Steele", "Noel Page",
    "Ari Flynn", "Lane Barrett", "Tatum York", "Jules Pratt", "Sasha Bloom",
    "Corey Vance", "Toby Hines", "Briar Kent", "Lennox Roy", "Zion Peak",
]
_FAKE_EMAIL = [
    "john.doe@example.com", "jane.smith@example.com", "peter.parker@example.com",
    "alex.morgan@example.com", "taylor.brown@example.com", "jordan.lee@example.com",
    "sam.wilson@example.com", "chris.blake@example.com", "morgan.reed@example.com",
    "casey.quinn@example.com", "riley.adams@example.com", "drew.campbell@example.com",
    "avery.hart@example.com", "jamie.stone@example.com", "skyler.west@example.com",
    "logan.grey@example.com", "dakota.mills@example.com", "reese.palmer@example.com",
    "hayden.brooks@example.com", "finley.shaw@example.com",
]
_FAKE_PHONE = [
    "+91 90000 00001", "+91 90000 00002", "+91 90000 00003",
    "+91 90000 00004", "+91 90000 00005", "+91 90000 00006",
    "+91 90000 00007", "+91 90000 00008", "+91 90000 00009",
    "+91 90000 00010", "+91 90000 00011", "+91 90000 00012",
    "+91 90000 00013", "+91 90000 00014", "+91 90000 00015",
    "+91 90000 00016", "+91 90000 00017", "+91 90000 00018",
    "+91 90000 00019", "+91 90000 00020",
]
_FAKE_COMPANY = [
    "Acme Corp Ltd.", "Globex Industries Limited", "Umbrella Holdings Private Limited",
    "Initech Solutions Limited", "Hooli Technologies Limited", "Pied Piper Inc.",
    "Stark Industries Limited", "Wayne Enterprises Limited", "Oscorp Limited",
    "Vance Refrigeration LLC", "Dunder Mifflin Private Limited", "Soylent Corp Limited",
    "Tyrell Corporation", "Cyberdyne Systems Limited", "Massive Dynamic Inc.",
    "Aperture Science LLC", "Wonka Industries Limited", "Weyland Corp Limited",
    "Bluth Company LLC", "Sterling Cooper Limited",
]
_FAKE_ADDRESS = [
    "123, MG Road, Sector 5, New Delhi – 110 001",
    "456, Park Street, Block B, Kolkata – 700 016",
    "789, Brigade Road, 2nd Floor, Bengaluru – 560 025",
    "101, Anna Salai, T Nagar, Chennai – 600 017",
    "202, FC Road, Shivajinagar, Pune – 411 005",
    "303, SG Highway, Thaltej, Ahmedabad – 380 054",
    "404, Link Road, Andheri West, Mumbai – 400 053",
    "505, Banjara Hills, Road No 12, Hyderabad – 500 034",
    "606, Residency Road, Richmond Town, Bengaluru – 560 025",
    "707, Civil Lines, Prayagraj – 211 001",
]
_FAKE_SSN = [
    "000-12-3456", "000-34-5678", "000-56-7890", "000-78-9012", "000-90-1234",
]
_FAKE_CC = [
    "4000 0000 0000 0000", "4111 0000 0000 0000", "4222 0000 0000 0000",
    "4333 0000 0000 0000", "4444 0000 0000 0000",
]
_FAKE_DOB = [
    "date of birth: 01/01/1990", "date of birth: 15/06/1985",
    "date of birth: 22/03/1978", "date of birth: 10/11/1992",
    "date of birth: 05/09/2000",
]
_FAKE_IP = [
    "10.0.0.1", "10.0.0.2", "10.0.0.3", "10.0.0.4", "10.0.0.5",
]

_FAKE_POOLS: dict[str, list[str]] = {
    "PERSON": _FAKE_PERSON,
    "EMAIL": _FAKE_EMAIL,
    "PHONE": _FAKE_PHONE,
    "COMPANY": _FAKE_COMPANY,
    "ADDRESS": _FAKE_ADDRESS,
    "SSN": _FAKE_SSN,
    "CREDIT_CARD": _FAKE_CC,
    "DOB": _FAKE_DOB,
    "IP_ADDRESS": _FAKE_IP,
}


class FakeFactory:
    def __init__(self):
        self.counters: dict[str, int] = {}
        self.mapping: dict[tuple[str, str], str] = {}

    def fake(self, pii_type: str, original: str) -> str:
        key = (pii_type, original.lower().strip())
        if key in self.mapping:
            return self.mapping[key]
        n = self.counters.get(pii_type, 0) + 1
        self.counters[pii_type] = n
        value = f"[{pii_type}_{n:03d}]"
        self.mapping[key] = value
        return value


def redact_text(text: str, factory: FakeFactory) -> tuple[str, list[tuple[str, str, str]]]:
    spans = detect_pii(text)
    replacements = []
    out = []
    cursor = 0
    for s in spans:
        original = text[s.start:s.end]
        fake = factory.fake(s.pii_type, original)
        out.append(text[cursor:s.start])
        out.append(fake)
        replacements.append((s.pii_type, original, fake))
        cursor = s.end
    out.append(text[cursor:])
    return ''.join(out), replacements


def docx_to_docx(input_docx: Path, output_docx: Path) -> tuple[int, int, dict[str, int]]:
    doc = Document(str(input_docx))
    factory = FakeFactory()
    category_counts: dict[str, int] = {}
    total_replacements = 0
    paragraph_count = len(doc.paragraphs)

    def process_p(p):
        nonlocal total_replacements
        if not p.text:
            return
        text = normalize(p.text)
        redacted, reps = redact_text(text, factory)
        if reps:
            total_replacements += len(reps)
            for typ, _, _ in reps:
                category_counts[typ] = category_counts.get(typ, 0) + 1
            p.text = redacted

    for p in doc.paragraphs:
        process_p(p)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    process_p(p)

    for section in doc.sections:
        for header in (section.header, getattr(section, 'first_page_header', None)):
            if header and hasattr(header, 'paragraphs'):
                for p in header.paragraphs:
                    process_p(p)
                for t in getattr(header, 'tables', []):
                    for row in t.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                process_p(p)
        for footer in (section.footer, getattr(section, 'first_page_footer', None)):
            if footer and hasattr(footer, 'paragraphs'):
                for p in footer.paragraphs:
                    process_p(p)
                for t in getattr(footer, 'tables', []):
                    for row in t.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                process_p(p)

    doc.save(str(output_docx))
    return paragraph_count, total_replacements, category_counts


def pdf_to_docx(pdf_path: Path, output_docx: Path) -> tuple[int, int, dict[str, int]]:
    reader = PdfReader(str(pdf_path))
    doc = Document()
    factory = FakeFactory()
    category_counts: dict[str, int] = {}
    total_replacements = 0
    page_count = 0

    for idx, page in enumerate(reader.pages, start=1):
        raw = page.extract_text() or ""
        text = normalize(raw)
        redacted, reps = redact_text(text, factory)
        total_replacements += len(reps)
        page_count += 1

        for typ, _, _ in reps:
            category_counts[typ] = category_counts.get(typ, 0) + 1

        p = doc.add_paragraph()
        run = p.add_run(f"Page {idx}")
        run.bold = True
        doc.add_paragraph(redacted)
        if idx != len(reader.pages):
            doc.add_page_break()

    doc.save(str(output_docx))
    return page_count, total_replacements, category_counts


def print_summary_table(input_file: Path, output_docx: Path, unit_label: str, unit_count: int, total_reps: int, category_counts: dict[str, int]) -> None:
    border = "=" * 62
    divider = "-" * 62
    print("\n" + border)
    print("                PII REDACTION TOOL — SUMMARY                ")
    print(border)
    print(f" Input File  : {input_file.name}")
    print(f" Output DOCX : {output_docx.name}")
    print(f" Scope       : {unit_count} {unit_label}")
    print(f" Total PII   : {total_reps} redaction replacements made")
    print(divider)
    print(f" {'PII Category':<25} | {'Redactions Made':<20}")
    print(divider)
    
    cat_names = {
        "PERSON": "Full Names",
        "COMPANY": "Company Names",
        "EMAIL": "Email Addresses",
        "ADDRESS": "Physical Addresses",
        "PHONE": "Phone Numbers",
        "SSN": "Social Security Numbers",
        "CREDIT_CARD": "Credit Card Numbers",
        "DOB": "Dates of Birth",
        "IP_ADDRESS": "IP Addresses",
    }
    
    for cat, name in cat_names.items():
        count = category_counts.get(cat, 0)
        print(f" {name:<25} | {count:<20}")
    print(border + "\n")


def main() -> None:
    parser = argparse.ArgumentParser(
        description="PII Redaction Tool — Redacts sensitive PII from PDF/DOCX into a clean DOCX"
    )
    parser.add_argument("input_file", type=Path, help="Input PDF or DOCX file path")
    parser.add_argument("output_docx", type=Path, nargs="?", default=None, help="Output DOCX file path (optional)")
    args = parser.parse_args()

    if not args.input_file.exists():
        raise FileNotFoundError(f"Input file not found: {args.input_file}")

    output_path = args.output_docx
    if output_path is None:
        output_path = args.input_file.with_name("KSH_PII_Redacted_RHP.docx")

    if args.input_file.suffix.lower() == ".docx":
        units, total_reps, cat_counts = docx_to_docx(args.input_file, output_path)
        print_summary_table(args.input_file, output_path, "paragraphs", units, total_reps, cat_counts)
    else:
        pages, total_reps, cat_counts = pdf_to_docx(args.input_file, output_path)
        print_summary_table(args.input_file, output_path, "pages", pages, total_reps, cat_counts)


if __name__ == "__main__":
    main()


